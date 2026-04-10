import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
import os
import pandas as pd
import time
import win32com.client as win32

# === 依赖库检查 ===
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    messagebox.showerror("缺少库", "请先安装 python-docx！\n运行: pip install python-docx")

# ==========================================
# 0. 模块中英文映射字典
# ==========================================
BLOCK_TYPE_MAP = {
    "HEATX": "换热器",
    "HEATER": "加热器/冷却器",
    "RADFRAC": "严格精馏塔",
    "MCOMPR": "多段压缩机",
    "FLASH2": "两相闪蒸罐",
    "FLASH3": "三相闪蒸罐",
    "PUMP": "泵",
    "COMPR": "压缩机",
    "RPLUG": "平推流反应器",
    "RSTOIC": "化学计量反应器"
}


# ==========================================
# 1. 核心提取逻辑
# ==========================================

def Hot_Stream(sim, Stream_Name):
    try:
        path = f"\\Data\\Streams\\{Stream_Name}\\Output\\HMX_FLOW\\MIXED"
        hot_node = sim.Tree.FindNode(path)
        return hot_node.Value * 0.0041868 if hot_node else 0
    except:
        return 0


def Get_Stream_Data(sim, Stream_Name):
    """提取基础数据：固定返回长度为8的列表"""
    try:
        node = sim.Tree.Elements("Data").Elements("Streams").Elements(Stream_Name).Elements("Output")
        temp_k = node.Elements("STR_MAIN").Elements("TEMP").Elements("MIXED").Value
        pres_atm = node.Elements("STR_MAIN").Elements("PRES").Elements("MIXED").Value

        data = [
            Stream_Name,  # 0. 流股名称
            (temp_k or 273.15) - 273.15,  # 1. 温度
            (pres_atm or 0) * 101.325,  # 2. 压力
            node.Elements("MASSVFRA").Value,  # 3. 气相分率
            node.Elements("RES_MOLEFLOW").Value,  # 4. 摩尔流量
            node.Elements("RES_MASSFLOW").Value,  # 5. 质量流量
            node.Elements("VOLFLMX2").Value * 0.06,  # 6. 体积流量
            Hot_Stream(sim, Stream_Name)  # 7. 焓值(kW)
        ]
        return data
    except:
        return [Stream_Name, 0, 0, 0, 0, 0, 0, 0]


# ==========================================
# 2. GUI 主程序
# ==========================================

class MrsJoneEnergyGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("MrsJone 能量衡算")
        self.root.geometry("950x850")

        # === 新增：彩蛋键盘监听缓冲区 ===
        self.key_buffer = ""
        self.root.bind("<Key>", self.check_easter_egg)
        # ==============================

        self.sim = None
        self.doc = None
        self.all_blocks = []

        self.aspen_file = tk.StringVar()
        self.save_dir = tk.StringVar()
        self.visibility = tk.BooleanVar(value=False)
        self.status = tk.StringVar(value="准备就绪")

        self.gen_word = tk.BooleanVar(value=True)
        self.gen_excel = tk.BooleanVar(value=True)
        self.sort_mode = tk.StringVar(value="default")

        self.export_options = {
            "流股名称": tk.BooleanVar(value=True),
            "温度(℃)": tk.BooleanVar(value=True),
            "压力(kPa)": tk.BooleanVar(value=True),
            "质量汽相分率": tk.BooleanVar(value=True),
            "摩尔流量(kmol/h)": tk.BooleanVar(value=True),
            "质量流量(kg/h)": tk.BooleanVar(value=True),
            "体积流量(m3/h)": tk.BooleanVar(value=True),
            "焓值(kW)": tk.BooleanVar(value=True)
        }

        self.create_widgets()
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    # === 新增：彩蛋触发函数 ===
    def check_easter_egg(self, event):
        if event.char:
            # 记录按下的按键并转为小写
            self.key_buffer += event.char.lower()
            # 保持缓冲区长度为 7 (mrsjone 的长度)
            self.key_buffer = self.key_buffer[-7:]

            # 检查是否匹配彩蛋口令
            if self.key_buffer == "mrsjone":
                messagebox.showinfo("🎉 发现彩蛋", "Hello, Aspener!")
                self.key_buffer = ""  # 触发后清空缓冲区，防止连续触发

    # ==========================

    def create_widgets(self):
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=10, pady=5)

        self.tab_main = ttk.Frame(self.notebook)
        self.tab_preview = ttk.Frame(self.notebook)

        self.notebook.add(self.tab_main, text="💻 控制面板")
        self.notebook.add(self.tab_preview, text="👁️‍🗨️ 实时能量平衡面板")

        self.setup_main_tab()
        self.setup_preview_tab()

        ttk.Label(self.root, textvariable=self.status, relief="sunken", anchor="w").pack(side="bottom", fill="x")

    def setup_main_tab(self):
        main = ttk.Frame(self.tab_main, padding=10)
        main.pack(fill="both", expand=True)

        group1 = ttk.LabelFrame(main, text="1. 环境连接", padding=10)
        group1.pack(fill="x", pady=5)

        f_frame = ttk.Frame(group1)
        f_frame.pack(fill="x")
        ttk.Entry(f_frame, textvariable=self.aspen_file, width=80).pack(side="left", padx=5)
        ttk.Button(f_frame, text="选择BKP文件", command=self.browse_file).pack(side="left")

        d_frame = ttk.Frame(group1)
        d_frame.pack(fill="x", pady=5)
        ttk.Entry(d_frame, textvariable=self.save_dir, width=80).pack(side="left", padx=5)
        ttk.Button(d_frame, text="选择输出目录", command=self.browse_folder).pack(side="left")
        ttk.Button(d_frame, text="打开输出目录", command=self.open_output_folder).pack(side="left", padx=10)

        conn_f = ttk.Frame(group1)
        conn_f.pack(fill="x", pady=5)
        ttk.Checkbutton(conn_f, text="显示 Aspen 界面", variable=self.visibility).pack(side="left", padx=5)

        self.status_canvas = tk.Canvas(conn_f, width=16, height=16, highlightthickness=0)
        self.status_canvas.pack(side="left", padx=(15, 2))
        self.indicator = self.status_canvas.create_oval(2, 2, 14, 14, fill="red", outline="gray")
        ttk.Label(conn_f, text="连接状态").pack(side="left", padx=(0, 15))

        self.conn_btn = ttk.Button(conn_f, text="连接 Aspen (读取模块)", command=self.connect_aspen)
        self.conn_btn.pack(side="left", padx=5)
        self.close_btn = ttk.Button(conn_f, text="断开 Aspen", command=self.close_aspen, state="disabled")
        self.close_btn.pack(side="left", padx=5)

        group2 = ttk.LabelFrame(main, text="2. 指定模块提取 (按住 Ctrl 可多选)", padding=10)
        group2.pack(fill="x", pady=5)

        list_frame = ttk.Frame(group2)
        list_frame.pack(fill="x")

        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side="right", fill="y")
        self.block_listbox = tk.Listbox(list_frame, selectmode="extended", height=5, yscrollcommand=scrollbar.set)
        self.block_listbox.pack(side="left", fill="x", expand=True)
        scrollbar.config(command=self.block_listbox.yview)

        btn_frame = ttk.Frame(group2)
        btn_frame.pack(fill="x", pady=2)
        ttk.Button(btn_frame, text="全选", command=self.select_all_blocks).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="清空选择", command=self.clear_selection).pack(side="left")

        group3 = ttk.LabelFrame(main, text="3. 导出控制与表格字段", padding=10)
        group3.pack(fill="x", pady=5)

        format_frame = ttk.Frame(group3)
        format_frame.pack(fill="x", pady=2)
        ttk.Label(format_frame, text="报告格式:").pack(side="left", padx=5)
        ttk.Checkbutton(format_frame, text="Word 汇总", variable=self.gen_word).pack(side="left", padx=5)
        ttk.Checkbutton(format_frame, text="Excel 明细", variable=self.gen_excel).pack(side="left")

        ttk.Label(format_frame, text=" |  Word 输出排序:").pack(side="left", padx=(15, 5))
        ttk.Radiobutton(format_frame, text="默认(按提取顺序)", variable=self.sort_mode, value="default").pack(
            side="left", padx=5)
        ttk.Radiobutton(format_frame, text="按模块类型(如换热器分组)", variable=self.sort_mode, value="by_type").pack(
            side="left", padx=5)

        ttk.Separator(group3, orient='horizontal').pack(fill='x', pady=8)

        ttk.Label(group3, text="表格提取字段 (表2流股明细):").pack(anchor="w", padx=5)
        checkbox_frame = ttk.Frame(group3)
        checkbox_frame.pack(fill="x", pady=2)
        col, row = 0, 0
        for text, var in self.export_options.items():
            ttk.Checkbutton(checkbox_frame, text=text, variable=var).grid(row=row, column=col, sticky="w", padx=15,
                                                                          pady=2)
            col += 1
            if col > 3:
                col = 0
                row += 1

        group4 = ttk.LabelFrame(main, text="4. 运行", padding=10)
        group4.pack(fill="x", pady=5)
        self.run_btn = ttk.Button(group4, text="开始提取选中模块", command=self.start_batch_run, state="disabled")
        self.run_btn.pack(fill="x", pady=5)
        self.progress = ttk.Progressbar(group4, mode="determinate")
        self.progress.pack(fill="x", pady=5)

        info_frame = ttk.Frame(main, padding=5)
        info_frame.pack(fill="x")

        log_tool_frame = ttk.Frame(info_frame)
        log_tool_frame.pack(fill="x")
        ttk.Label(log_tool_frame, text="沧州交通学院校训: 进德修业，知行合一", foreground="blue",
                  font=("微软雅黑", 9, "bold")).pack(side="left")
        ttk.Button(log_tool_frame, text="🧹 清空日志", command=self.clear_log).pack(side="right")

        self.log_text = scrolledtext.ScrolledText(main, height=8, state='disabled', font=("Consolas", 9))
        self.log_text.pack(fill="both", expand=True, pady=5)

    def setup_preview_tab(self):
        preview_frame = ttk.Frame(self.tab_preview, padding=10)
        preview_frame.pack(fill="both", expand=True)

        tool_frame = ttk.Frame(preview_frame)
        tool_frame.pack(fill="x", pady=(0, 5))
        ttk.Label(tool_frame, text="此表格实时展示各个模块的能量衡算结果。误差(Error)较大时将标红显示。",
                  foreground="gray").pack(side="left")
        ttk.Button(tool_frame, text="清空表格", command=self.clear_treeview).pack(side="right")

        tree_frame = ttk.Frame(preview_frame)
        tree_frame.pack(fill="both", expand=True)

        scroll_y = ttk.Scrollbar(tree_frame, orient="vertical")
        scroll_y.pack(side="right", fill="y")
        scroll_x = ttk.Scrollbar(tree_frame, orient="horizontal")
        scroll_x.pack(side="bottom", fill="x")

        columns = ("Block", "Type", "W", "Q", "Hin", "Hout", "Error")
        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings", yscrollcommand=scroll_y.set,
                                 xscrollcommand=scroll_x.set)

        self.tree.tag_configure("error_high", foreground="red", font=("微软雅黑", 9, "bold"))
        self.tree.tag_configure("normal", foreground="black")

        scroll_y.config(command=self.tree.yview)
        scroll_x.config(command=self.tree.xview)

        headings = ["所属模块", "模块类型", "轴功 W (kW)", "热负荷 Q (kW)", "进料总焓 Hin (kW)", "出料总焓 Hout (kW)",
                    "平衡误差 Error (kW)"]
        widths = [120, 150, 110, 110, 130, 130, 140]

        for col, head, w in zip(columns, headings, widths):
            self.tree.heading(col, text=head)
            self.tree.column(col, width=w, anchor="center")
        self.tree.pack(side="left", fill="both", expand=True)

    def set_indicator(self, color):
        self.status_canvas.itemconfig(self.indicator, fill=color)
        self.root.update()

    def log(self, msg, level="INFO"):
        self.log_text.config(state='normal')
        tag = "red" if level == "ERROR" else ("green" if level == "SUCCESS" else "black")
        self.log_text.tag_config("red", foreground="red")
        self.log_text.tag_config("green", foreground="darkgreen")
        self.log_text.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] {msg}\n", tag)
        self.log_text.see(tk.END)
        self.log_text.config(state='disabled')
        self.root.update()

    def clear_log(self):
        self.log_text.config(state='normal')
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state='disabled')
        self.log("日志已清空。")

    def clear_treeview(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

    def select_all_blocks(self):
        self.block_listbox.select_set(0, tk.END)

    def clear_selection(self):
        self.block_listbox.selection_clear(0, tk.END)

    def connect_aspen(self):
        if not self.aspen_file.get():
            messagebox.showwarning("提示", "请先选择BKP文件")
            return
        try:
            self.set_indicator("yellow")
            self.log("正在初始化 Aspen 进程 (时间较长请耐心等待)...")
            self.conn_btn.config(state="disabled")
            self.root.update()

            f = self.aspen_file.get()
            self.sim = win32.gencache.EnsureDispatch("Apwn.Document")
            self.sim.InitFromArchive2(os.path.abspath(f))
            self.sim.Visible = self.visibility.get()

            self.log("正在运行模拟以刷新数据...")
            self.sim.Tree.Elements('Data').Elements("Setup").Elements("Global").Elements("Input").Elements(
                "INSET").Value = "MET"
            self.sim.Run()

            self.log("正在读取流程中的模块列表...")
            blocks = self.sim.Tree.Elements("Data").Elements("Blocks").Elements
            self.all_blocks = [blk.Name for blk in blocks]

            self.block_listbox.delete(0, tk.END)
            for b_name in self.all_blocks:
                self.block_listbox.insert(tk.END, b_name)
            self.select_all_blocks()

            self.run_btn.config(state="normal")
            self.close_btn.config(state="normal")
            self.set_indicator("green")
            self.log(f"Aspen 连接成功！共找到 {len(self.all_blocks)} 个模块。", "SUCCESS")
            self.status.set("Aspen 已连接 - 请选择要提取的模块")
        except Exception as e:
            self.set_indicator("red")
            self.conn_btn.config(state="normal")
            self.log(f"连接失败: {e}", "ERROR")

    def close_aspen(self):
        if self.sim:
            self.sim = None
            self.log("Aspen 连接已断开。")
            self.block_listbox.delete(0, tk.END)
        self.run_btn.config(state="disabled")
        self.close_btn.config(state="disabled")
        self.conn_btn.config(state="normal")
        self.set_indicator("red")
        self.status.set("已断开")

    def start_batch_run(self):
        selected_indices = self.block_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("提示", "请在模块列表中至少选择一个要提取的模块！")
            return

        if not (self.gen_word.get() or self.gen_excel.get()):
            messagebox.showwarning("提示", "请至少勾选一种要生成的报告类型！")
            return

        save_path = self.save_dir.get()
        if not save_path:
            messagebox.showwarning("提示", "请选择输出保存目录！")
            return
        os.chdir(save_path)

        self.clear_treeview()

        if self.gen_word.get():
            self.doc = Document()

            main_heading = self.doc.add_heading('', level=0)
            run = main_heading.add_run('Aspen Plus 选定模块能量衡算报告')
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.color.rgb = RGBColor(0, 0, 0)

            self.doc.add_paragraph(f"导出时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            self.doc.add_paragraph("有问题可以加联系方式: wenxiaoshuo12138")
            self.doc.add_paragraph("可以加群讨论: 562721026")

        selected_blocks = [self.block_listbox.get(idx) for idx in selected_indices]
        self.progress["maximum"] = len(selected_blocks)
        self.log(f"=== 开始提取任务 (共选定 {len(selected_blocks)} 个模块) ===")

        self.set_indicator("yellow")
        self.run_btn.config(state="disabled")

        if self.sort_mode.get() == "by_type":
            self.log("正在分析模块类型并重新排序...", "INFO")
            block_dict = {}
            for b_name in selected_blocks:
                try:
                    b_type_raw = str(
                        self.sim.Tree.Elements("Data").Elements("Blocks").Elements(b_name).AttributeValue(6)).upper()
                    if b_type_raw not in block_dict:
                        block_dict[b_type_raw] = []
                    block_dict[b_type_raw].append(b_name)
                except:
                    pass

            count = 0
            for b_type_raw, b_names in block_dict.items():
                if self.gen_word.get():
                    p = self.doc.add_paragraph()
                    type_zh = BLOCK_TYPE_MAP.get(b_type_raw, b_type_raw)
                    run = p.add_run(f"\n【{type_zh} ({b_type_raw}) 设备组汇总】")
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 51, 153)
                for b_name in b_names:
                    count += 1
                    self.progress["value"] = count
                    self.process_single_block(b_name)
                    self.root.update()
        else:
            for i, b_name in enumerate(selected_blocks):
                self.progress["value"] = i + 1
                self.process_single_block(b_name)
                self.root.update()

        if self.gen_word.get():
            fname = f"能量衡算报告_{int(time.time())}.docx"
            self.doc.save(fname)
            self.log(f"【提取完成】汇总Word已保存: {fname}", "SUCCESS")

        if self.gen_excel.get():
            self.log("【提取完成】选定模块的 Excel 数据已生成。", "SUCCESS")

        self.set_indicator("green")
        self.run_btn.config(state="normal")
        self.notebook.select(self.tab_preview)

        if messagebox.askyesno("全部完成",
                               "提取任务已完成！\n数据已在[实时能量平衡面板]展示。\n是否立即打开输出文件夹查看报告？"):
            self.open_output_folder()

    def process_single_block(self, b_name):
        try:
            blk = self.sim.Tree.Elements("Data").Elements("Blocks").Elements(b_name)
            b_type = str(blk.AttributeValue(6)).upper()
            port = blk.Elements("Ports")

            b_type_zh = BLOCK_TYPE_MAP.get(b_type, b_type)

            selected_headers = [name for name, var in self.export_options.items() if var.get()]
            if not selected_headers:
                self.log(f"跳过 {b_name}：未勾选任何表格字段", "ERROR")
                return

            table_cols = []
            header = pd.Series(selected_headers, name=" ")
            table_cols.append(header)

            W, Q, Hin, Hout = 0, 0, 0, 0
            is_valid = False

            if b_type in ["HEATX", "RADFRAC", "MCOMPR", "FLASH2", "FLASH3", "PUMP", "COMPR", "RPLUG", "HEATER",
                          "RSTOIC"]:
                is_valid = True
                ta = [i.Name for i in port.Elements]

                for tag in ta:
                    for i in [n.Name for n in port.Elements(tag).Elements]:
                        if [n.Name for n in port.Elements(tag).Elements] is None: break
                        raw_data = Get_Stream_Data(self.sim, i)
                        direction = "IN" if "IN" in tag else "OUT"

                        if direction == "IN":
                            Hin += raw_data[7]
                        else:
                            Hout += raw_data[7]

                        filtered_data = [raw_data[idx] for idx, (key, var) in enumerate(self.export_options.items()) if
                                         var.get()]
                        s = pd.Series(filtered_data)
                        s.name = f"In ({i})" if direction == "IN" else f"Out ({i})"
                        table_cols.append(s)

                if b_type in ["HEATER", "RPLUG"]:
                    Q = self.sim.Tree.FindNode(f'\\Data\\Blocks\\{b_name}\\Output\\QCALC').Value * 0.0041868
                elif b_type == "RADFRAC":
                    reb = self.sim.Tree.FindNode(f'\\Data\\Blocks\\{b_name}\\Output\\REB_DUTY').Value * 0.0041868
                    cond = self.sim.Tree.FindNode(f'\\Data\\Blocks\\{b_name}\\Output\\COND_DUTY').Value * 0.0041868
                    Q = reb + cond
                elif b_type in ["FLASH2", "FLASH3"]:
                    Q = self.sim.Tree.FindNode(f'\\Data\\Blocks\\{b_name}\\Output\\QCALC').Value * 0.0041868
                elif b_type in ["PUMP", "COMPR"]:
                    W = self.sim.Tree.FindNode(f'\\Data\\Blocks\\{b_name}\\Output\\WNET').Value
                elif b_type == "MCOMPR":
                    W = self.sim.Tree.FindNode(f'\\Data\\Blocks\\{b_name}\\Output\\QCALC2').Value
                    Q = self.sim.Tree.FindNode(f'\\Data\\Blocks\\{b_name}\\Output\\DUTY_OUT').Value * 0.0041868

            if is_valid:
                error = Hin - Hout + Q + W

                def fmt(val):
                    return f"{val:.4f}" if isinstance(val, (int, float)) else str(val)

                tree_values = (b_name, b_type_zh, fmt(W), fmt(Q), fmt(Hin), fmt(Hout), fmt(error))

                row_tag = "error_high" if abs(error) > 0.01 else "normal"
                self.tree.insert("", "end", values=tree_values, tags=(row_tag,))

                if self.gen_excel.get():
                    self.export_individual_excel(b_name, table_cols, W, Q, Hin, Hout)
                if self.gen_word.get():
                    self.append_to_word(b_name, b_type_zh, b_type, table_cols, W, Q, Hin, Hout)
                self.log(f"模块 {b_name} ({b_type_zh}) 提取成功", "SUCCESS")
            else:
                self.log(f"模块 {b_name} 类型 [{b_type}] 暂不支持或非衡算模块，已跳过。")

        except Exception as e:
            self.log(f"模块 {b_name} 处理异常: {e}", "ERROR")

    def export_individual_excel(self, b_name, table_cols, W, Q, Hin, Hout):
        try:
            df1 = pd.concat(table_cols, axis=1)
            df2 = pd.DataFrame([W, Q, Hin, Hout, Hin - Hout + Q + W], index=["W", "Q", "Hin", "Hout", "Error"]).T
            df3 = pd.DataFrame(["热负荷", W, Q], index=["项目", "W(kW)", "Q(kW)"]).T
            with pd.ExcelWriter(f"{b_name}-能量衡算.xlsx") as writer:
                df3.to_excel(writer, sheet_name='表1-负荷表', index=False)
                df1.to_excel(writer, sheet_name='表2-流股明细', index=False)
                df2.to_excel(writer, sheet_name='表3-平衡验证', index=False)
        except Exception as e:
            self.log(f"导出 Excel 失败: {e}", "ERROR")

    def append_to_word(self, b_name, b_type_zh, b_type_raw, table_cols, W, Q, Hin, Hout):
        doc = self.doc
        doc.add_page_break()

        p_heading = doc.add_paragraph()
        run = p_heading.add_run(f'模块: {b_name} ({b_type_zh} - {b_type_raw})')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True

        df1 = pd.DataFrame(["热负荷/功", W, Q], index=["项目", "W (kW)", "Q (kW)"]).T
        df2 = pd.concat(table_cols, axis=1)
        df3 = pd.DataFrame([W, Q, Hin, Hout, Hin - Hout + Q + W],
                           index=["W(kW)", "Q(kW)", "Hin(kW)", "Hout(kW)", "Error"]).T

        for title, df in [("表1：负荷数据", df1), ("表2：流股数据明细", df2), ("表3：能量平衡验证", df3)]:
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t_run = p_title.add_run(title)
            t_run.font.size = Pt(10.5)
            t_run.font.name = '黑体'
            t_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            table.style = 'Table Grid'
            for j, col in enumerate(df.columns):
                table.cell(0, j).text = str(col)
                table.cell(0, j).paragraphs[0].runs[0].font.bold = True
            for i, row in enumerate(df.values):
                for j, val in enumerate(row):
                    table.cell(i + 1, j).text = f"{val:.4f}" if isinstance(val, (float, int)) else str(val)

    def browse_file(self):
        f = filedialog.askopenfilename(filetypes=[("Aspen Backup Files", "*.bkp"), ("All Files", "*.*")])
        if f: self.aspen_file.set(f)

    def browse_folder(self):
        d = filedialog.askdirectory()
        if d: self.save_dir.set(d)

    def open_output_folder(self):
        path = self.save_dir.get()
        if path and os.path.exists(path):
            try:
                os.startfile(path)
                self.log("已为您打开输出文件夹。")
            except Exception as e:
                self.log(f"无法打开文件夹: {e}", "ERROR")
        else:
            messagebox.showwarning("提示", "请先选择一个有效的输出目录！")

    def on_closing(self):
        if self.sim: self.close_aspen()
        self.root.destroy()


if __name__ == "__main__":
    root = tk.Tk()
    app = MrsJoneEnergyGUI(root)
    root.mainloop()